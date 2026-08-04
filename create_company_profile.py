"""
יצירת פרופיל חברה - Twenty2Jobs / HR22 Group
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===========================
# הגדרות עיצוב
# ===========================
PRIMARY_COLOR   = RGBColor(0, 168, 168)   # טורקיז
DARK_COLOR      = RGBColor(0, 102, 102)   # ירוק כהה
WHITE           = RGBColor(255, 255, 255)
BLACK           = RGBColor(30, 30, 30)
GRAY            = RGBColor(100, 100, 100)
LIGHT_BG        = RGBColor(240, 250, 250)
GOLD            = RGBColor(180, 130, 0)

FONT_NAME = "Arial"

def set_rtl(paragraph):
    """הגדרת כיוון RTL לפסקה"""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

def set_rtl_table(table):
    """הגדרת RTL לטבלה"""
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    bidiVisual = OxmlElement('w:bidiVisual')
    tblPr.append(bidiVisual)

def add_run_rtl(paragraph, text, bold=False, size=12, color=None, italic=False):
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # RTL for run
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)
    return run

def add_heading(doc, text, level=1, color=None, size=None, center=False):
    p = doc.add_paragraph()
    set_rtl(p)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if level == 1:
        sz = size or 24
        clr = color or PRIMARY_COLOR
        bold = True
    elif level == 2:
        sz = size or 18
        clr = color or DARK_COLOR
        bold = True
    else:
        sz = size or 14
        clr = color or BLACK
        bold = True
    add_run_rtl(p, text, bold=bold, size=sz, color=clr)
    return p

def add_paragraph(doc, text, size=11, color=None, bold=False, center=False, italic=False):
    p = doc.add_paragraph()
    set_rtl(p)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clr = color or BLACK
    add_run_rtl(p, text, bold=bold, size=size, color=clr, italic=italic)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.left_indent = Cm(1)
    add_run_rtl(p, "◆  " + text, size=size, color=BLACK)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_divider(doc, color=None):
    p = doc.add_paragraph()
    set_rtl(p)
    clr = color or PRIMARY_COLOR
    run = p.add_run("━" * 55)
    run.font.color.rgb = clr
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)

def add_section_box(doc, title, color=None):
    """כותרת סעיף עם רקע צבעוני"""
    if color is None:
        hex_color = '00A8A8'
    elif isinstance(color, str):
        hex_color = color
    else:
        # RGBColor is a tuple (r, g, b)
        hex_color = '%02X%02X%02X' % (color[0], color[1], color[2])
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # רקע
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run("  " + title + "  ")
    run.font.name = FONT_NAME
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = WHITE
    return p

def create_logo_text(doc):
    """לוגו טקסטואלי בראש הדף"""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_rtl_table(table)
    
    # עמודה שמאל - בס"ד
    cell_l = table.cell(0, 0)
    cell_l.width = Cm(4)
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_l.add_run('בס"ד')
    r.font.name = FONT_NAME
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = DARK_COLOR

    # עמודה מרכז - לוגו
    cell_c = table.cell(0, 1)
    cell_c.width = Cm(10)
    p_c = cell_c.paragraphs[0]
    p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # רקע כחול-ירוק ללוגו
    pPr = p_c._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '00A8A8')
    pPr.append(shd)
    
    r22 = p_c.add_run('  22  ')
    r22.font.name = FONT_NAME
    r22.font.size = Pt(36)
    r22.font.bold = True
    r22.font.color.rgb = WHITE
    
    rj = p_c.add_run('JOBS')
    rj.font.name = FONT_NAME
    rj.font.size = Pt(13)
    rj.font.bold = True
    rj.font.color.rgb = RGBColor(200, 255, 255)
    
    p_sub = cell_c.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:val'), 'clear')
    shd2.set(qn('w:color'), 'auto')
    shd2.set(qn('w:fill'), '006666')
    p_sub._p.get_or_add_pPr().append(shd2)
    rs = p_sub.add_run('Twenty2Jobs | HR22 Group')
    rs.font.name = FONT_NAME
    rs.font.size = Pt(9)
    rs.font.bold = False
    rs.font.color.rgb = RGBColor(200, 255, 255)

    # עמודה ימין - ריק
    cell_r = table.cell(0, 2)
    cell_r.width = Cm(4)
    
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()
    
    # הגדרות דף
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    
    # RTL לכל המסמך
    doc.core_properties.language = 'he-IL'
    settings = doc.settings.element
    bidi = OxmlElement('w:bidi')
    settings.append(bidi)

    # ============================
    # כותרת עמוד - לוגו + בס"ד
    # ============================
    create_logo_text(doc)

    # כותרת ראשית
    p_title = doc.add_paragraph()
    set_rtl(p_title)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(6)
    r = p_title.add_run('פרופיל חברה')
    r.font.name = FONT_NAME
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = DARK_COLOR

    p_sub = doc.add_paragraph()
    set_rtl(p_sub)
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = p_sub.add_run('חברת ההשמה המובילה בישראל | שנת 2026')
    rs.font.name = FONT_NAME
    rs.font.size = Pt(13)
    rs.font.color.rgb = GRAY
    rs.italic = True
    p_sub.paragraph_format.space_after = Pt(4)

    add_divider(doc)

    # ============================
    # 1. אודות החברה
    # ============================
    add_section_box(doc, '◀  אודות החברה  ▶')
    
    add_paragraph(doc,
        'קבוצת 22 טו-גדר בע"מ – Twenty2Jobs היא חברת השמה ומיון עובדים מובילה בישראל, '
        'המתמחה בחיבור בין מועמדים מוכשרים לבין מעסיקים מהמגזר הפרטי והציבורי. '
        'אנו פועלים בשיטת הצלחה בלבד – גייסת, שילמת. אם לא גויס עובד – לא משלמים.',
        size=12)

    add_paragraph(doc,
        'הפלטפורמה שלנו, הנגישה בכתובת www.hr22group.com, מציעה מאות משרות פעילות '
        'בתחומי הלוגיסטיקה, הרכב, הייטק, מכירות, פיננסים, שירות לקוחות ועוד, '
        'עם כלים מתקדמים המבוססים על בינה מלאכותית.',
        size=12)

    doc.add_paragraph()

    # ============================
    # 2. שיטת ההצלחה
    # ============================
    add_section_box(doc, '💎  מודל גייסת – שילמת  💎', color='006666')

    add_paragraph(doc,
        'אנו עובדים אך ורק בשיטת הצלחה מלאה:', size=12, bold=True)

    add_bullet(doc, 'פרסמנו את המשרה? לא משלמים.')
    add_bullet(doc, 'ראיינו מועמדים? לא משלמים.')
    add_bullet(doc, 'גייסנו עובד שהתקבל ועבד? רק אז מגיעה עמלת ההצלחה.')
    add_bullet(doc, 'לא מצאנו את האדם המתאים? אין עלות.')
    add_bullet(doc, 'שקיפות מלאה בכל שלב בתהליך.')

    add_paragraph(doc,
        'גישה זו מבטיחה שהאינטרסים שלנו מיושרים לחלוטין עם האינטרסים שלכם – '
        'ההצלחה שלכם היא ההצלחה שלנו.',
        size=11, color=GRAY, italic=True)

    doc.add_paragraph()

    # ============================
    # 3. תחומי התמחות
    # ============================
    add_section_box(doc, '🚀  תחומי התמחות')

    areas = [
        ('🚛 לוגיסטיקה ומחסנאות', 'מנהלי לוגיסטיקה, מלגזנים, מלקטים, מחסנאים, בקרי סחורה, רפרנטים'),
        ('🚗 עולם הרכב', 'מכונאים, חשמלאי רכב, יועצי שירות, מנהלי סניפים, אנשי מכירות'),
        ('💼 מכירות ושיווק', 'נציגי מכירות, מנהלי תיקי לקוחות, מנהלי צוות מכירות'),
        ('🏦 פיננסים ובנקאות', 'פקידי בנק, יועצי השקעות, רואי חשבון, כלכלנים'),
        ('📞 שירות לקוחות', 'נציגי שירות, ראשי צוות מוקד, מנהלי שירות'),
        ('💻 הייטק ומחשבים', 'מפתחים, DevOps, מנהלי מוצר, QA, UX/UI'),
        ('🏭 ייצור ותעשייה', 'עובדי ייצור, מנהלי מפעל, מהנדסים, מפעילי מכונות'),
        ('🏨 מלונאות ומזון', 'שפים, מנהלי מסעדות, אנשי קבלה, מאורגנים'),
    ]

    for title_a, desc in areas:
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(title_a + ' — ')
        r1.font.name = FONT_NAME
        r1.font.size = Pt(11)
        r1.font.bold = True
        r1.font.color.rgb = DARK_COLOR
        r2 = p.add_run(desc)
        r2.font.name = FONT_NAME
        r2.font.size = Pt(11)
        r2.font.color.rgb = BLACK

    doc.add_paragraph()

    # ============================
    # 4. לקוחות מובילים
    # ============================
    add_section_box(doc, '🏆  לקוחות מובילים שבחרו בנו')

    clients = [
        'Toyota Israel', 'Zeekr', 'Geely', 'Lexus Israel',
        'Union Group', 'בנק מזרחי טפחות', 'קישרי ים', 'סלע לוגיסטיקה',
        'מרלוג', 'Upright', 'GAC', 'בוטיק הפיתה', 'MRG', 'ועוד רבים...'
    ]

    # טבלת לקוחות
    rows = [clients[i:i+4] for i in range(0, len(clients), 4)]
    tbl = doc.add_table(rows=len(rows), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_rtl_table(tbl)

    for ri, row_data in enumerate(rows):
        for ci, client in enumerate(row_data):
            cell = tbl.cell(ri, ci)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'E0F5F5')
            pPr.append(shd)
            r = p.add_run(client)
            r.font.name = FONT_NAME
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = DARK_COLOR

    doc.add_paragraph()

    # ============================
    # 5. מערכת ה-CRM - Twenty2CRM
    # ============================
    add_section_box(doc, '💻  מערכת Twenty2CRM – הטכנולוגיה שלנו')

    add_paragraph(doc,
        'פיתחנו בעצמנו מערכת CRM מתקדמת המאפשרת ניהול מלא ואוטומטי של כל '
        'מחזור חיי הגיוס – ממועמד ראשון ועד סגירת משרה. המערכת בנויה על Next.js, '
        'Prisma ובינה מלאכותית (Google Gemini) ופועלת בענן.',
        size=12, bold=False)

    features = [
        ('🧠 ניתוח קורות חיים באמצעות AI',
         'Gemini AI סורק ומנתח קורות חיים אוטומטית – מחלץ נתונים, מגלה ניסיון '
         'ומייצר פרופיל מועמד מלא תוך שניות, כולל זיהוי עברית ואנגלית.'),
        ('🎯 מנוע התאמה חכמה',
         'אלגוריתם מתקדם שמתאים מועמדים למשרות לפי מיקום, ניסיון, שכר, כישורים '
         'ותגיות – ומדרג את המועמדים לפי רלוונטיות.'),
        ('📧 סריקת מיילים אוטומטית',
         'חיבור ל-Gmail API – המערכת סורקת תיבת דואר נכנס אוטומטית, מזהה '
         'קורות חיים ומכניסה מועמדים חדשים ישירות למסד הנתונים.'),
        ('📊 ניהול מועמדים מלא',
         'מעקב אחר כל מועמד בכל שלב – הגשה, ראיון, הצעת עבודה, קבלה. '
         'היסטוריה מלאה, הערות, תגיות ויצירת קשר ישירה.'),
        ('📋 ניהול משרות ומעסיקים',
         'ממשק עשיר לניהול מאות משרות עם תגיות, מיקום, שכר, שעות ותנאים. '
         'סנכרון אוטומטי לאתר hr22group.com.'),
        ('📱 WhatsApp & SMS',
         'שליחת הודעות מותאמות אישית למועמדים ישירות מהמערכת – לא צריך לצאת.'),
        ('📅 ניהול ראיונות',
         'יומן ראיונות, תזכורות אוטומטיות, ניהול Zoom/Teams/פגישה פיזית.'),
        ('📈 דוחות ואנליטיקה',
         'לוח מחוונים עם נתוני גיוס בזמן אמת – כמה מועמדים, משרות, ראיונות '
         'ויחס סגירה.'),
        ('☁️ גיבוי אוטומטי ל-Google Drive',
         'כל הנתונים מגובים אוטומטית בענן. אפס סיכון לאיבוד מידע.'),
        ('🌐 סנכרון עם אתר Twenty2Jobs',
         'משרות מתעדכנות אוטומטית באתר hr22group.com ברגע שנפתחות או נסגרות.'),
    ]

    for icon_title, desc in features:
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        
        r1 = p.add_run(icon_title + '\n')
        r1.font.name = FONT_NAME
        r1.font.size = Pt(11)
        r1.font.bold = True
        r1.font.color.rgb = PRIMARY_COLOR
        
        r2 = p.add_run('     ' + desc)
        r2.font.name = FONT_NAME
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = GRAY

    doc.add_paragraph()

    # ============================
    # 6. הצוותים שלנו
    # ============================
    add_section_box(doc, '👥  הצוותים שלנו')

    teams = [
        ('צוות גיוס לוגיסטיקה ותעשייה',
         'מגייסים מנוסים המתמחים בתפקידי מחסן, לוגיסטיקה, ייצור ותעשייה. '
         'מכירים היטב את שוק הלוגיסטיקה הישראלי ואת הצרכים הייחודיים של חברות '
         'כמו סלע לוגיסטיקה, מרלוג, קישרי ים ועוד.'),
        ('צוות גיוס רכב ומכירות',
         'מגייסים בעלי רקע בעולם הרכב והמכירות, עם קשרים ישירים ל-Toyota, Zeekr, '
         'Geely ו-Lexus Israel. מתמחים בגיוס מכונאים, חשמלאים ואנשי מכירות.'),
        ('צוות גיוס פיננסים ובנקאות',
         'מגייסים מנוסים עם הבנה מעמיקה בשוק הפיננסי ובדרישות בנק מזרחי טפחות '
         'ומוסדות פיננסיים נוספים.'),
        ('צוות טכנולוגיה ופיתוח',
         'מפתחי Full-Stack הבונים ומתחזקים את מערכת Twenty2CRM ואת פלטפורמת '
         'hr22group.com. עבודה עם Next.js, TypeScript, Prisma, Gemini AI ו-Railway.'),
        ('צוות שיווק ומדיה',
         'מטפלים בנוכחות הדיגיטלית, SEO, ניהול מדיה חברתית, קידום משרות '
         'ובניית המותג Twenty2Jobs בישראל.'),
        ('צוות תמיכה ושירות לקוחות',
         'זמינים לכל מעסיק ומועמד. מספקים מענה מהיר, מקצועי ואישי – '
         'כי אצלנו כל אחד חשוב.'),
    ]

    for team_name, team_desc in teams:
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run('▸ ' + team_name)
        r1.font.name = FONT_NAME
        r1.font.size = Pt(12)
        r1.font.bold = True
        r1.font.color.rgb = DARK_COLOR
        
        p2 = doc.add_paragraph()
        set_rtl(p2)
        p2.paragraph_format.left_indent = Cm(1)
        p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run(team_desc)
        r2.font.name = FONT_NAME
        r2.font.size = Pt(11)
        r2.font.color.rgb = BLACK

    doc.add_paragraph()

    # ============================
    # 7. כלים נוספים לדרך
    # ============================
    add_section_box(doc, '🛠️  כלים חינמיים שאנו מציעים')

    tools = [
        '🧮 מחשבון שכר נטו/ברוטו – החישוב המדויק ביותר בישראל',
        '⏰ שעון נוכחות דיגיטלי – מעקב שעות לעובדים',
        '⚖️ מדריך חוקי עבודה – כל מה שצריך לדעת',
        '📱 אפליקציה לנייד – גישה מכל מקום',
        '🤖 אביגדור – סוכן AI חכם שממליץ על משרות לפי קורות החיים',
        '📰 מגזין קריירה – טיפים, מדריכים ועדכונים שוטפים',
        '🎁 תוכנית "המלץ על חבר" – תגמול על כל המלצה מוצלחת',
    ]
    for t in tools:
        add_bullet(doc, t, size=11)

    doc.add_paragraph()

    # ============================
    # 8. למה לבחור בנו
    # ============================
    add_section_box(doc, '⭐  למה לבחור ב-Twenty2Jobs?')

    reasons = [
        'שיטת גייסת-שילמת – אפס סיכון פיננסי למעסיק',
        '246+ משרות פעילות בכל רגע נתון',
        'מאגר של אלפי מועמדים מסורקים ומוכנים',
        'AI מתקדם לניתוח קורות חיים ב-30 שניות',
        'מנוע התאמה חכמה – מועמד הנכון, בזמן הנכון',
        'לקוחות מהחברות הגדולות ביותר בישראל',
        'זמינות גבוהה – מענה מהיר ואישי',
        'סנכרון מלא בין ה-CRM לאתר הדרושים הציבורי',
        'גיבוי ואבטחת מידע ברמה הגבוהה ביותר',
        'פיתוח מתמיד של כלים וטכנולוגיות חדשות',
    ]
    for r in reasons:
        add_bullet(doc, r)

    doc.add_paragraph()

    # ============================
    # 9. פרטי התקשרות
    # ============================
    add_section_box(doc, '📞  יצירת קשר')

    contact_info = [
        ('🌐 אתר', 'www.hr22group.com'),
        ('📧 דוא"ל', 'hr@hr22group.com'),
        ('📱 טלפון', 'זמין דרך האתר'),
        ('📍 שם משפטי', 'קבוצת 22 טו-גדר בע"מ'),
        ('🔗 LinkedIn', 'חיפוש "Twenty2Jobs" בלינקדאין'),
    ]

    for label, value in contact_info:
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(label + ': ')
        r1.font.name = FONT_NAME
        r1.font.size = Pt(11)
        r1.font.bold = True
        r1.font.color.rgb = DARK_COLOR
        r2 = p.add_run(value)
        r2.font.name = FONT_NAME
        r2.font.size = Pt(11)
        r2.font.color.rgb = BLACK

    doc.add_paragraph()
    add_divider(doc)

    # סיום
    p_end = doc.add_paragraph()
    set_rtl(p_end)
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run('© 2026 קבוצת 22 טו-גדר בע"מ | Twenty2Jobs | כל הזכויות שמורות')
    r_end.font.name = FONT_NAME
    r_end.font.size = Pt(9)
    r_end.font.color.rgb = GRAY
    r_end.italic = True

    p_bsd = doc.add_paragraph()
    set_rtl(p_bsd)
    p_bsd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_bsd = p_bsd.add_run('בס"ד')
    r_bsd.font.name = FONT_NAME
    r_bsd.font.size = Pt(11)
    r_bsd.font.bold = True
    r_bsd.font.color.rgb = DARK_COLOR

    # שמירה
    output_path = r'C:\Users\22ged\OneDrive\Desktop\TWENTY2CRM\פרופיל_חברה_Twenty2Jobs_2026.docx'
    doc.save(output_path)
    print(f'✅ הקובץ נשמר: {output_path}')
    return output_path


if __name__ == '__main__':
    path = build_document()
    print('🎉 פרופיל החברה נוצר בהצלחה!')
    print(f'📄 נתיב: {path}')
